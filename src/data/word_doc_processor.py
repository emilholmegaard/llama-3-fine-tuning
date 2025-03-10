"""
Enhanced Word document processor for Llama 3.3 fine-tuning.

This module extracts and processes content from Word documents (.doc and .docx),
preserving folder structure for categorization, with support for extracting
tables and images, and outputting structured JSON files.
"""

import os
import logging
import json
import base64
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union, Any
from dataclasses import dataclass, field

import docx
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.exceptions import PackageNotFoundError

import mammoth
from bs4 import BeautifulSoup
from tqdm import tqdm
import olefile  # For handling .doc (legacy) files

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


@dataclass
class ProcessingOptions:
    """Options for Word document processing."""
    
    recursive: bool = True
    preserve_structure: bool = True
    extract_images: bool = False
    extract_tables: bool = True
    extract_metadata: bool = True
    extract_headers_footers: bool = True
    min_text_length: int = 50
    max_documents: int = -1
    file_extensions: List[str] = field(default_factory=lambda: [".docx", ".doc"])
    image_format: str = "base64"  # Options: base64, filename
    output_format: str = "json"  # Options: json, jsonl, plain_text
    include_raw_html: bool = False  # Whether to include raw HTML in output
    paragraph_separator: str = "\n\n"  # Separator between paragraphs


class WordDocProcessor:
    """Enhanced Word document processor for fine-tuning."""

    def __init__(
        self,
        input_dir: str,
        output_dir: str,
        options: Optional[ProcessingOptions] = None,
    ):
        """
        Initialize the Word document processor.

        Args:
            input_dir: Directory containing Word documents
            output_dir: Directory to save processed documents
            options: Processing options (if None, default options are used)
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.options = options or ProcessingOptions()

        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Set up image directory if needed
        if self.options.extract_images and self.options.image_format == "filename":
            self.images_dir = self.output_dir / "images"
            os.makedirs(self.images_dir, exist_ok=True)

    def find_documents(self) -> List[Path]:
        """
        Find all Word documents in the input directory.

        Returns:
            List of paths to Word documents
        """
        files = []
        search_pattern = "**/*" if self.options.recursive else "*"

        for ext in self.options.file_extensions:
            pattern = search_pattern + ext
            found_files = list(self.input_dir.glob(pattern))
            files.extend(found_files)

        logger.info(f"Found {len(files)} Word documents")
        
        if self.options.max_documents > 0:
            files = files[:self.options.max_documents]
            logger.info(f"Processing first {self.options.max_documents} documents")
            
        return files

    def get_output_path(self, input_path: Path) -> Path:
        """
        Get the output path for a processed document.

        Args:
            input_path: Path to the input document

        Returns:
            Path to the output file
        """
        suffix = f".{self.options.output_format.lower()}"
        
        if self.options.preserve_structure:
            # Preserve the folder structure
            rel_path = input_path.relative_to(self.input_dir)
            output_path = self.output_dir / rel_path.with_suffix(suffix)
            os.makedirs(output_path.parent, exist_ok=True)
        else:
            # Flatten the structure
            output_path = self.output_dir / f"{input_path.stem}{suffix}"

        return output_path

    def _extract_images_from_docx(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """
        Extract images from a docx Document object.
        
        Args:
            doc: The python-docx Document object
            
        Returns:
            List of image dictionaries with metadata
        """
        images = []
        
        if not self.options.extract_images:
            return images
            
        try:
            # Get document part relationships
            rels = doc.part.rels
            
            for rel_id, rel in rels.items():
                # Check if relationship is an image
                if rel.reltype == RT.IMAGE:
                    try:
                        image_part = rel.target_part
                        image_filename = image_part.partname.split('/')[-1]
                        content_type = image_part.content_type
                        
                        image_data = {
                            "filename": image_filename,
                            "content_type": content_type,
                        }
                        
                        # Extract image data based on chosen format
                        if self.options.image_format == "base64":
                            img_bytes = image_part.blob
                            img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                            image_data["data"] = img_b64
                            image_data["encoding"] = "base64"
                        else:  # filename
                            # Save image to file
                            rel_doc_path = self.output_dir.relative_to(self.output_dir.parent)
                            img_path = self.images_dir / rel_doc_path / image_filename
                            os.makedirs(img_path.parent, exist_ok=True)
                            
                            with open(img_path, 'wb') as f:
                                f.write(image_part.blob)
                                
                            image_data["path"] = str(img_path.relative_to(self.output_dir))
                            
                        images.append(image_data)
                    except Exception as e:
                        logger.warning(f"Error extracting image {rel_id}: {e}")
        except Exception as e:
            logger.warning(f"Error accessing document relationships: {e}")
            
        return images

    def _extract_tables_from_docx(self, doc: DocxDocument) -> List[List[List[str]]]:
        """
        Extract tables from a docx Document object.
        
        Args:
            doc: The python-docx Document object
            
        Returns:
            List of tables, each represented as a list of rows of cells
        """
        tables = []
        
        if not self.options.extract_tables:
            return tables
            
        try:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
        except Exception as e:
            logger.warning(f"Error extracting tables: {e}")
            
        return tables

    def _extract_headers_footers_from_docx(self, doc: DocxDocument) -> Dict[str, List[str]]:
        """
        Extract headers and footers from a docx Document object.
        
        Args:
            doc: The python-docx Document object
            
        Returns:
            Dictionary with headers and footers content
        """
        headers_footers = {
            "headers": [],
            "footers": []
        }
        
        if not self.options.extract_headers_footers:
            return headers_footers
            
        try:
            # Extract headers
            for section in doc.sections:
                try:
                    if section.header.is_linked_to_previous:
                        continue
                        
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            headers_footers["headers"].append(paragraph.text.strip())
                except Exception as e:
                    logger.debug(f"Error extracting header: {e}")
                    
                try:
                    if section.footer.is_linked_to_previous:
                        continue
                        
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            headers_footers["footers"].append(paragraph.text.strip())
                except Exception as e:
                    logger.debug(f"Error extracting footer: {e}")
        except Exception as e:
            logger.warning(f"Error extracting headers and footers: {e}")
            
        return headers_footers

    def extract_text_with_python_docx(self, doc_path: Path) -> Dict:
        """
        Extract text and metadata from a Word document using python-docx.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = docx.Document(doc_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables if requested
            tables = self._extract_tables_from_docx(doc) if self.options.extract_tables else []
            
            # Extract images if requested
            images = self._extract_images_from_docx(doc) if self.options.extract_images else []
            
            # Extract headers and footers if requested
            headers_footers = self._extract_headers_footers_from_docx(doc) if self.options.extract_headers_footers else {"headers": [], "footers": []}
            
            # Extract document properties
            properties = {}
            if self.options.extract_metadata:
                try:
                    core_props = doc.core_properties
                    properties = {
                        "title": core_props.title or "",
                        "author": core_props.author or "",
                        "created": str(core_props.created) if core_props.created else "",
                        "modified": str(core_props.modified) if core_props.modified else "",
                        "last_modified_by": core_props.last_modified_by or "",
                    }
                except Exception as e:
                    logger.warning(f"Error extracting metadata: {e}")
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "images": images,
                "headers": headers_footers["headers"],
                "footers": headers_footers["footers"],
                "properties": properties,
                "full_text": self.options.paragraph_separator.join(paragraphs),
            }
            
        except PackageNotFoundError:
            # File is not a docx file, might be a legacy .doc file
            return self.extract_text_from_doc(doc_path)
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path} using python-docx: {e}")
            return None

    def extract_text_from_doc(self, doc_path: Path) -> Dict:
        """
        Extract text from a legacy .doc file using olefile.
        
        Args:
            doc_path: Path to the .doc document
            
        Returns:
            Dictionary with extracted text
        """
        try:
            if not olefile.isOleFile(doc_path):
                logger.warning(f"{doc_path} is not a valid OLE file")
                return None
                
            with olefile.OleFile(doc_path) as ole:
                # Try to extract text from the WordDocument stream
                if ole.exists('WordDocument'):
                    # For .doc files, we can only extract basic text
                    # Full .doc parsing is very complex; using a simplified approach
                    try:
                        # Use the internal metadata if available
                        metadata = {}
                        if ole.exists('\005SummaryInformation'):
                            si_stream = ole.openstream('\005SummaryInformation')
                            metadata = {"from_summary_info": True}
                            
                        # Try to extract raw text - this is not a complete solution
                        # but provides basic text extraction
                        text_chunks = []
                        wordstream = ole.openstream('WordDocument')
                        word_data = wordstream.read()
                        
                        # Look for text chunks (very simplified approach)
                        for chunk in word_data.split(b'\x00\x00\x00'):
                            if len(chunk) > 10:  # Filter out tiny chunks
                                try:
                                    # Try to decode as text
                                    text = chunk.replace(b'\x00', b'').decode('utf-8', errors='ignore')
                                    if len(text.strip()) > 0:
                                        text_chunks.append(text.strip())
                                except:
                                    pass
                                    
                        # Return simplified data structure for .doc files
                        return {
                            "paragraphs": text_chunks,
                            "tables": [],
                            "images": [],
                            "headers": [],
                            "footers": [],
                            "properties": metadata,
                            "full_text": self.options.paragraph_separator.join(text_chunks),
                            "doc_format": "legacy_doc"
                        }
                    except Exception as e:
                        logger.warning(f"Error parsing WordDocument stream: {e}")
                        
            # If we reach here, try mammoth as a backup for .doc files
            return self.extract_text_with_mammoth(doc_path)
            
        except Exception as e:
            logger.error(f"Error extracting text from .doc file {doc_path}: {e}")
            return None

    def extract_text_with_mammoth(self, doc_path: Path) -> Dict:
        """
        Extract text and metadata from a Word document using mammoth.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            with open(doc_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(html, "lxml")
                
                # Extract paragraphs
                paragraphs = [p.get_text() for p in soup.find_all("p") if p.get_text().strip()]
                
                # Extract tables if requested
                tables = []
                if self.options.extract_tables:
                    for table in soup.find_all("table"):
                        table_data = []
                        for row in table.find_all("tr"):
                            row_data = [cell.get_text().strip() for cell in row.find_all(["td", "th"])]
                            table_data.append(row_data)
                        tables.append(table_data)
                
                # Extract images if requested 
                images = []
                if self.options.extract_images:
                    for idx, img in enumerate(soup.find_all("img")):
                        images.append({
                            "alt_text": img.get("alt", ""),
                            "index": idx,
                            "format": "unknown"
                        })
                
                # Extract document properties (limited info available)
                properties = {
                    "messages": [message for message in result.messages],
                }
                
                result_data = {
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "images": images,
                    "headers": [],  # Mammoth doesn't extract headers/footers
                    "footers": [],
                    "properties": properties,
                    "full_text": self.options.paragraph_separator.join(paragraphs),
                }
                
                # Include raw HTML if requested
                if self.options.include_raw_html:
                    result_data["raw_html"] = html
                    
                return result_data
                
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path} using mammoth: {e}")
            return None

    def process_document(self, doc_path: Path) -> Optional[Dict]:
        """
        Process a single Word document.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with processed content or None if processing failed
        """
        try:
            # Determine file type based on extension
            is_docx = doc_path.suffix.lower() == ".docx"
            
            # Try python-docx first for .docx files, or mammoth as fallback
            if is_docx:
                content = self.extract_text_with_python_docx(doc_path)
            else:
                # For .doc files, try our .doc extractor first, which falls back to mammoth
                content = self.extract_text_from_doc(doc_path)
                
            # If primary extraction failed, try the alternative method
            if content is None:
                content = self.extract_text_with_mammoth(doc_path)
                
            if content is None:
                logger.warning(f"Failed to extract content from {doc_path}")
                return None
                
            # Check if the document has enough text
            if len(content["full_text"]) < self.options.min_text_length:
                logger.warning(f"Document {doc_path} has insufficient text content, skipping")
                return None
                
            # Add document path information
            content["metadata"] = {
                "filename": doc_path.name,
                "path": str(doc_path),
                "relative_path": str(doc_path.relative_to(self.input_dir)),
                "category": str(doc_path.parent.relative_to(self.input_dir)) if self.options.preserve_structure else "",
                "file_type": "docx" if is_docx else "doc",
                "extraction_method": "python-docx" if is_docx else "doc_extractor",
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing document {doc_path}: {e}")
            return None

    def write_output(self, content: Dict, output_path: Path) -> bool:
        """
        Write processed content to the output file.
        
        Args:
            content: Processed document content
            output_path: Path to write the output file
            
        Returns:
            True if writing was successful, False otherwise
        """
        try:
            os.makedirs(output_path.parent, exist_ok=True)
            
            if self.options.output_format.lower() == "json":
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(content, f, ensure_ascii=False, indent=2)
            
            elif self.options.output_format.lower() == "jsonl":
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(json.dumps(content, ensure_ascii=False))
                    
            elif self.options.output_format.lower() == "plain_text":
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content["full_text"])
                    
            else:
                logger.error(f"Unsupported output format: {self.options.output_format}")
                return False
                
            return True
                
        except Exception as e:
            logger.error(f"Error writing output to {output_path}: {e}")
            return False

    def process_all(self) -> Tuple[int, int]:
        """
        Process all Word documents in the input directory.

        Returns:
            Tuple of (number of documents processed, number of documents succeeded)
        """
        documents = self.find_documents()
        succeeded = 0
        
        for doc_path in tqdm(documents, desc="Processing documents"):
            content = self.process_document(doc_path)
            
            if content:
                output_path = self.get_output_path(doc_path)
                
                if self.write_output(content, output_path):
                    succeeded += 1
                
        logger.info(f"Processed {len(documents)} documents, {succeeded} succeeded")
        return len(documents), succeeded


def main():
    """Main function for command-line execution."""
    import argparse

    parser = argparse.ArgumentParser(description="Enhanced Word document processor for Llama fine-tuning")
    
    # Required arguments
    parser.add_argument("--input_dir", required=True, help="Directory containing Word documents")
    parser.add_argument("--output_dir", required=True, help="Directory to save processed documents")
    
    # Processing options
    parser.add_argument("--recursive", action="store_true", default=True, help="Process documents in subdirectories")
    parser.add_argument("--preserve_structure", action="store_true", default=True, 
                       help="Maintain folder structure in output")
    parser.add_argument("--extract_images", action="store_true", help="Extract images from documents")
    parser.add_argument("--extract_tables", action="store_true", default=True, help="Extract tables from documents")
    parser.add_argument("--min_text_length", type=int, default=50, help="Minimum text length to keep")
    
    args = parser.parse_args()
    
    # Create processor and process documents
    processor = WordDocProcessor(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        options=ProcessingOptions(
            recursive=args.recursive,
            preserve_structure=args.preserve_structure,
            extract_images=args.extract_images,
            extract_tables=args.extract_tables,
            min_text_length=args.min_text_length,
        ),
    )
    
    processor.process_all()


if __name__ == "__main__":
    main()
