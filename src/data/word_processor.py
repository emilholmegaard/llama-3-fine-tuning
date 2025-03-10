"""
Word document processor for Llama 3.3 fine-tuning.

This module extracts and processes content from Word documents, preserving
folder structure for categorization if needed.
"""

import os
import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import docx
import mammoth
from bs4 import BeautifulSoup
from tqdm import tqdm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class WordProcessor:
    """Process Word documents for fine-tuning."""

    def __init__(
        self,
        input_dir: str,
        output_dir: str,
        recursive: bool = True,
        preserve_structure: bool = True,
        extract_images: bool = False,
        extract_tables: bool = True,
        min_text_length: int = 50,
        max_documents: int = -1,
        file_extensions: List[str] = None,
    ):
        """
        Initialize the Word document processor.

        Args:
            input_dir: Directory containing Word documents
            output_dir: Directory to save processed documents
            recursive: Process documents in subdirectories
            preserve_structure: Maintain folder structure in output
            extract_images: Extract image descriptions
            extract_tables: Extract tables from documents
            min_text_length: Minimum text length to keep
            max_documents: Maximum number of documents to process (-1 for all)
            file_extensions: List of file extensions to process
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.recursive = recursive
        self.preserve_structure = preserve_structure
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.min_text_length = min_text_length
        self.max_documents = max_documents
        self.file_extensions = file_extensions or [".docx", ".doc"]

        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)

    def find_documents(self) -> List[Path]:
        """
        Find all Word documents in the input directory.

        Returns:
            List of paths to Word documents
        """
        files = []
        search_pattern = "**/*" if self.recursive else "*"

        for ext in self.file_extensions:
            pattern = search_pattern + ext
            found_files = list(self.input_dir.glob(pattern))
            files.extend(found_files)

        logger.info(f"Found {len(files)} Word documents")
        
        if self.max_documents > 0:
            files = files[:self.max_documents]
            logger.info(f"Processing first {self.max_documents} documents")
            
        return files

    def get_output_path(self, input_path: Path) -> Path:
        """
        Get the output path for a processed document.

        Args:
            input_path: Path to the input document

        Returns:
            Path to the output file
        """
        if self.preserve_structure:
            # Preserve the folder structure
            rel_path = input_path.relative_to(self.input_dir)
            output_path = self.output_dir / rel_path.with_suffix(".json")
            os.makedirs(output_path.parent, exist_ok=True)
        else:
            # Flatten the structure
            output_path = self.output_dir / f"{input_path.stem}.json"

        return output_path

    def extract_text_with_python_docx(self, doc_path: Path) -> Dict:
        """
        Extract text and metadata from a Word document using python-docx.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = docx.Document(doc_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables if requested
            tables = []
            if self.extract_tables:
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append(table_data)
            
            # Extract document properties
            properties = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "last_modified_by": doc.core_properties.last_modified_by or "",
            }
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "properties": properties,
                "full_text": "\n\n".join(paragraphs),
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path} using python-docx: {e}")
            return None

    def extract_text_with_mammoth(self, doc_path: Path) -> Dict:
        """
        Extract text and metadata from a Word document using mammoth.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            with open(doc_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(html, "lxml")
                
                # Extract paragraphs
                paragraphs = [p.get_text() for p in soup.find_all("p") if p.get_text().strip()]
                
                # Extract tables if requested
                tables = []
                if self.extract_tables:
                    for table in soup.find_all("table"):
                        table_data = []
                        for row in table.find_all("tr"):
                            row_data = [cell.get_text() for cell in row.find_all(["td", "th"])]
                            table_data.append(row_data)
                        tables.append(table_data)
                
                # Extract document properties (limited info available)
                properties = {
                    "messages": [message for message in result.messages],
                }
                
                return {
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "properties": properties,
                    "full_text": "\n\n".join(paragraphs),
                }
                
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path} using mammoth: {e}")
            return None

    def process_document(self, doc_path: Path) -> Optional[Dict]:
        """
        Process a single Word document.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary with processed content or None if processing failed
        """
        try:
            # Try python-docx first
            content = self.extract_text_with_python_docx(doc_path)
            
            # If python-docx fails, try mammoth
            if content is None:
                content = self.extract_text_with_mammoth(doc_path)
                
            if content is None:
                logger.warning(f"Failed to extract content from {doc_path}")
                return None
                
            # Check if the document has enough text
            if len(content["full_text"]) < self.min_text_length:
                logger.warning(f"Document {doc_path} has insufficient text content, skipping")
                return None
                
            # Add document path information
            content["metadata"] = {
                "filename": doc_path.name,
                "path": str(doc_path),
                "relative_path": str(doc_path.relative_to(self.input_dir)),
                "category": str(doc_path.parent.relative_to(self.input_dir)) if self.preserve_structure else "",
            }
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing document {doc_path}: {e}")
            return None

    def process_all(self) -> Tuple[int, int]:
        """
        Process all Word documents in the input directory.

        Returns:
            Tuple of (number of documents processed, number of documents succeeded)
        """
        documents = self.find_documents()
        succeeded = 0
        
        for doc_path in tqdm(documents, desc="Processing documents"):
            content = self.process_document(doc_path)
            
            if content:
                output_path = self.get_output_path(doc_path)
                os.makedirs(output_path.parent, exist_ok=True)
                
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(content, f, ensure_ascii=False, indent=2)
                    
                succeeded += 1
                
        logger.info(f"Processed {len(documents)} documents, {succeeded} succeeded")
        return len(documents), succeeded


def main():
    """Main function for command-line execution."""
    import argparse

    parser = argparse.ArgumentParser(description="Process Word documents for Llama fine-tuning")
    parser.add_argument("--input_dir", required=True, help="Directory containing Word documents")
    parser.add_argument("--output_dir", required=True, help="Directory to save processed documents")
    parser.add_argument("--recursive", action="store_true", help="Process documents in subdirectories")
    parser.add_argument("--preserve_structure", action="store_true", help="Maintain folder structure in output")
    parser.add_argument("--extract_images", action="store_true", help="Extract image descriptions")
    parser.add_argument("--extract_tables", action="store_true", help="Extract tables from documents")
    parser.add_argument("--min_text_length", type=int, default=50, help="Minimum text length to keep")
    parser.add_argument("--max_documents", type=int, default=-1, help="Maximum number of documents to process")
    parser.add_argument("--file_extensions", nargs="+", default=[".docx", ".doc"], help="File extensions to process")

    args = parser.parse_args()
    
    processor = WordProcessor(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        recursive=args.recursive,
        preserve_structure=args.preserve_structure,
        extract_images=args.extract_images,
        extract_tables=args.extract_tables,
        min_text_length=args.min_text_length,
        max_documents=args.max_documents,
        file_extensions=args.file_extensions,
    )
    
    processor.process_all()


if __name__ == "__main__":
    main()
